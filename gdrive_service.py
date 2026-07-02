"""
gdrive_service.py
-----------------
Google Drive integration – uploads DOCX files and returns a view URL.
The file is kept as DOCX (no conversion) so formatting is 100% preserved.

Setup (one-time):
  1. Go to https://console.cloud.google.com/
  2. Create a project → Enable "Google Drive API"
  3. Create OAuth 2.0 credentials (Desktop app)
  4. Download → save as credentials.json next to this file
"""

import os

SCOPES = ['https://www.googleapis.com/auth/drive']
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CREDENTIALS_FILE = os.path.join(BASE_DIR, 'credentials.json')
TOKEN_FILE = os.path.join(BASE_DIR, 'token.json')
# Target folder in user's Drive: "מכתבים - מחולל מכתבים"
DRIVE_FOLDER_ID    = '1Y8sf_BTFGysOGQDdXO74Jey2-0Y1IR59'
# Template file: "תבנית מכתב - ח״כ יונתן משריקי.docx"
DRIVE_TEMPLATE_ID  = '1GV5b3-vdyCZ_vO_Y4QJC8b9u2NnlLZjg'


def has_client_secrets() -> bool:
    return os.path.exists(CREDENTIALS_FILE)


def _load_creds():
    if not os.path.exists(TOKEN_FILE):
        return None
    try:
        from google.oauth2.credentials import Credentials
        from google.auth.transport.requests import Request
        creds = Credentials.from_authorized_user_file(TOKEN_FILE, SCOPES)
        if creds.expired and creds.refresh_token:
            creds.refresh(Request())
            with open(TOKEN_FILE, 'w', encoding='utf-8') as f:
                f.write(creds.to_json())
        return creds if creds.valid else None
    except Exception:
        return None


def has_valid_token() -> bool:
    return _load_creds() is not None


# Flow objects keyed by state — needed to preserve PKCE code_verifier
_pending_flows: dict = {}


def get_auth_url(redirect_uri: str) -> str:
    from google_auth_oauthlib.flow import Flow
    flow = Flow.from_client_secrets_file(
        CREDENTIALS_FILE, scopes=SCOPES, redirect_uri=redirect_uri
    )
    auth_url, state = flow.authorization_url(
        access_type='offline', prompt='consent'
    )
    _pending_flows[state] = flow
    return auth_url


def exchange_code(code: str, state: str, redirect_uri: str) -> None:
    flow = _pending_flows.pop(state, None)
    if flow is None:
        # fallback without PKCE — may work for non-PKCE flows
        from google_auth_oauthlib.flow import Flow
        flow = Flow.from_client_secrets_file(
            CREDENTIALS_FILE, scopes=SCOPES, redirect_uri=redirect_uri
        )
    flow.fetch_token(code=code)
    with open(TOKEN_FILE, 'w', encoding='utf-8') as f:
        f.write(flow.credentials.to_json())


def _get_target_folder(service) -> str:
    """Return the target folder ID, verifying it exists."""
    try:
        service.files().get(fileId=DRIVE_FOLDER_ID, fields='id').execute()
        return DRIVE_FOLDER_ID
    except Exception:
        # Fallback: create a folder in root
        folder = service.files().create(
            body={'name': 'מכתבים - מחולל מכתבים',
                  'mimeType': 'application/vnd.google-apps.folder'},
            fields='id',
        ).execute()
        return folder['id']


_TEMPLATE_PATH    = os.path.join(BASE_DIR, 'הכנסת.docx')
_LOGO_DOCX_PATH   = os.path.join(BASE_DIR, 'לוגו.docx')
_LOGO_IMAGE_KEY   = 'word/media/image1.jpeg'   # image inside לוגו.docx
_LOGO_IMG_W       = 1268   # px
_LOGO_IMG_H       = 1793   # px
_HEADER_CROP_TOP  = 360    # px: content ends ~313px; keep with padding
_FOOTER_START_PX  = 1689   # px: contact line
# Original anchor extent in לוגו.docx (EMU)
_LOGO_CX          = 7438292   # 20.66 cm
_LOGO_CY_FULL     = 10518924  # 29.22 cm
_LOGO_CY_HEADER   = int(_LOGO_CY_FULL * _HEADER_CROP_TOP / _LOGO_IMG_H)   # ~5.87 cm
_LOGO_CY_FOOTER   = int(_LOGO_CY_FULL * (_LOGO_IMG_H - _FOOTER_START_PX) / _LOGO_IMG_H)
_PAGE_WIDTH_TWIPS = 11906
_MARGIN_TWIPS     = 1800


def _get_logo_raw() -> bytes:
    """Return raw JPEG from לוגו.docx."""
    import zipfile
    with zipfile.ZipFile(_LOGO_DOCX_PATH) as z:
        return z.read(_LOGO_IMAGE_KEY)


def _get_image_strip(raw_jpeg: bytes, y0: int, y1: int) -> bytes:
    """Return a horizontal strip of a JPEG image as bytes."""
    from PIL import Image
    import io
    img = Image.open(io.BytesIO(raw_jpeg))
    strip = img.crop((0, y0, img.width, y1))
    buf = io.BytesIO()
    strip.save(buf, format='JPEG', quality=95)
    return buf.getvalue()


def _add_fullwidth_image_para(doc, img_bytes: bytes, after_element=None, before_element=None):
    """Add a paragraph containing an inline image that spans the full page width
    (using negative indents to extend beyond the text margins)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph
    from docx.shared import Twips, Emu
    from PIL import Image
    import io

    # Build the paragraph XML
    p_el = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')

    # Justify center
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)

    # Negative indents so the paragraph spans the full page width
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'),  str(-_MARGIN_TWIPS))
    ind.set(qn('w:right'), str(-_MARGIN_TWIPS))
    pPr.append(ind)

    # No space before/after
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'),  '0')
    pPr.append(spacing)

    p_el.append(pPr)

    # Insert into body
    body = doc.element.body
    if before_element is not None:
        idx = list(body).index(before_element)
        body.insert(idx, p_el)
    elif after_element is not None:
        idx = list(body).index(after_element) + 1
        body.insert(idx, p_el)
    else:
        body.insert(0, p_el)

    # Add the picture at full page width
    para_obj = Paragraph(p_el, doc)
    run = para_obj.add_run()
    img_io = io.BytesIO(img_bytes)
    # Width in EMU: page_width_twips / 1440 inches * 914400 EMU/inch
    page_emu = int(_PAGE_WIDTH_TWIPS / 1440 * 914400)
    run.add_picture(img_io, width=Emu(page_emu))
    return p_el


def _strip_headers_from_docx(src_path: str, dst_path: str) -> None:
    """Build a Drive-friendly DOCX using לוגו.docx as the header/footer source.

    - Grafts לוגו.docx header XML (with CROPPED image, correct extent) onto
      the generated letter so the floating letterhead no longer overlaps text.
    - Appends the footer contact-line strip inline before sectPr.
    - Substitutes David → David Libre (available in Google Drive).
    """
    import zipfile, re, io

    raw_logo = _get_logo_raw()
    header_bytes = _get_image_strip(raw_logo, 0, _HEADER_CROP_TOP)
    footer_bytes  = _get_image_strip(raw_logo, _FOOTER_START_PX, _LOGO_IMG_H)

    # ── Read source ZIP files ─────────────────────────────────────────
    with zipfile.ZipFile(src_path, 'r') as z:
        out = {n: z.read(n) for n in z.namelist()}

    with zipfile.ZipFile(_LOGO_DOCX_PATH, 'r') as z:
        logo = {n: z.read(n) for n in z.namelist()}

    # ── Graft header/footer XML from לוגו.docx ────────────────────────
    for key in logo:
        if re.match(r'word/(header|footer)\d*\.xml$', key):
            out[key] = logo[key]
        if re.match(r'word/_rels/(header|footer)\d*\.xml\.rels$', key):
            out[key] = logo[key]

    # ── Inject cropped images ─────────────────────────────────────────
    HKEY = 'word/media/logo_header.jpeg'
    FKEY = 'word/media/logo_footer.jpeg'
    out[HKEY] = header_bytes
    out[FKEY]  = footer_bytes

    # Fix header2.xml.rels → point to cropped header image
    h2rels = 'word/_rels/header2.xml.rels'
    if h2rels in out:
        x = out[h2rels].decode('utf-8')
        x = re.sub(r'Target="media/[^"]+"', 'Target="media/logo_header.jpeg"', x)
        out[h2rels] = x.encode('utf-8')

    # Fix header2.xml extent cy (full → cropped) + posV → page top
    h2xml = 'word/header2.xml'
    if h2xml in out:
        x = out[h2xml].decode('utf-8')
        # Update extent cy
        x = re.sub(
            r'(<wp:extent\s+cx="\d+"\s+cy=")(\d+)(")',
            lambda m: m.group(1) + str(_LOGO_CY_HEADER) + m.group(3),
            x
        )
        # Set posV to page, offset 0
        x = re.sub(
            r'(<wp:positionV\b[^>]*relativeFrom=")[^"]*(")',
            r'\g<1>page\g<2>', x
        )
        x = re.sub(
            r'(<wp:positionV\b[^>]*>.*?<wp:posOffset>)-?\d+(</wp:posOffset>)',
            r'\g<1>0\g<2>', x, flags=re.DOTALL
        )
        # Also fix extent in effectExtent / spPr if present
        x = re.sub(
            r'(<a:ext\s+cx="\d+"\s+cy=")(\d+)(")',
            lambda m: m.group(1) + str(_LOGO_CY_HEADER) + m.group(3),
            x
        )
        out[h2xml] = x.encode('utf-8')

    # ── Fix document.xml ──────────────────────────────────────────────
    doc_key = 'word/document.xml'
    if doc_key in out:
        xml = out[doc_key].decode('utf-8')

        # Font: David → David Libre
        xml = re.sub(r'(w:(?:ascii|hAnsi|cs))="David"', r'\1="David Libre"', xml)

        # Remove leading blank spacers from body
        xml = re.sub(
            r'(<w:body>)((?:\s*<w:p\b(?:(?!</w:p>).)*?<\/w:p>)*?)(<w:p\b)',
            lambda m: m.group(1) + re.sub(
                r'\s*<w:p\b(?:(?!</w:p>).)*?<\/w:p>',
                lambda pm: pm.group(0) if re.search(r'<w:t\b', pm.group(0)) else '',
                m.group(2)
            ) + m.group(3),
            xml, count=1, flags=re.DOTALL
        )

        # Append footer image para before sectPr
        footer_para = _build_inline_image_para_xml(
            'logo_footer', 'rIdLogoFooter', _LOGO_CX, _LOGO_CY_FOOTER
        )
        xml = xml.replace('<w:sectPr', footer_para + '<w:sectPr', 1)

        # Remove old headerReference / footerReference (from הכנסת.docx)
        xml = re.sub(r'<w:headerReference[^/]*/>', '', xml)
        xml = re.sub(r'<w:footerReference[^/]*/>', '', xml)

        # Inject לוגו.docx sectPr header/footer refs (insert before </w:sectPr>)
        logo_doc = logo.get('word/document.xml', b'').decode('utf-8')
        logo_refs = re.findall(r'<w:(?:header|footer)Reference\b[^/]*/>', logo_doc)
        if logo_refs:
            xml = xml.replace('</w:sectPr>', ''.join(logo_refs) + '</w:sectPr>', 1)

        out[doc_key] = xml.encode('utf-8')

    # ── Fix document.xml.rels ─────────────────────────────────────────
    doc_rels = 'word/_rels/document.xml.rels'
    if doc_rels in out:
        rels = out[doc_rels].decode('utf-8')
        # Drop old header/footer rels from letter template
        rels = re.sub(r'<Relationship\b[^>]*(header|footer)[^>]*/>\s*',
                      '', rels, flags=re.IGNORECASE)
        # Add footer image rel
        rels = rels.replace(
            '</Relationships>',
            '<Relationship Id="rIdLogoFooter" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
            'Target="media/logo_footer.jpeg"/>'
            '</Relationships>'
        )
        # Copy header/footer rels from לוגו.docx document rels
        logo_rels_xml = logo.get('word/_rels/document.xml.rels', b'').decode('utf-8')
        logo_hf_rels = re.findall(
            r'<Relationship\b[^>]*(header|footer)[^>]*/>', logo_rels_xml, flags=re.IGNORECASE
        )
        if logo_hf_rels:
            rels = rels.replace(
                '</Relationships>',
                '\n'.join(logo_hf_rels) + '</Relationships>'
            )
        out[doc_rels] = rels.encode('utf-8')

    # ── Write output ──────────────────────────────────────────────────
    with zipfile.ZipFile(dst_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in out.items():
            zout.writestr(name, data)


def _build_inline_image_para_xml(name: str, rel_id: str, cx: int, cy: int) -> str:
    """Return a <w:p> XML string with a full-page-width inline image."""
    return (
        f'<w:p>'
        f'<w:pPr><w:jc w:val="center"/>'
        f'<w:ind w:left="-{_MARGIN_TWIPS}" w:right="-{_MARGIN_TWIPS}"/>'
        f'<w:spacing w:before="0" w:after="0"/></w:pPr>'
        f'<w:r><w:rPr/><w:drawing>'
        f'<wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        f' distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="9998" name="{name}"/>'
        f'<wp:cNvGraphicFramePr>'
        f'<a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        f' noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        f'<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        f'<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:nvPicPr><pic:cNvPr id="9998" name="{name}"/>'
        f'<pic:cNvPicPr><a:picLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        f' noChangeAspect="1"/></pic:cNvPicPr></pic:nvPicPr>'
        f'<pic:blipFill>'
        f'<a:blip xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        f' r:embed="{rel_id}"'
        f' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>'
        f'<a:stretch xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        f'<a:fillRect/></a:stretch></pic:blipFill>'
        f'<pic:spPr>'
        f'<a:xfrm xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        f'<a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        f'<a:prstGeom xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        f' prst="rect"><a:avLst/></a:prstGeom>'
        f'</pic:spPr></pic:pic></a:graphicData></a:graphic>'
        f'</wp:inline></w:drawing></w:r></w:p>'
    )


def _download_drive_file(file_id: str, dest_path: str) -> None:
    """Download a Drive file by ID to a local path."""
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseDownload
    import io
    creds = _load_creds()
    service = build('drive', 'v3', credentials=creds)
    request = service.files().get_media(fileId=file_id)
    with io.FileIO(dest_path, 'wb') as fh:
        dl = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            _, done = dl.next_chunk()


def _upload_raw(file_path: str, filename: str) -> str:
    """Upload a DOCX to Drive without any header manipulation."""
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload
    creds = _load_creds()
    service = build('drive', 'v3', credentials=creds)
    folder_id = _get_target_folder(service)
    mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    media = MediaFileUpload(file_path, mimetype=mimetype, resumable=False)
    result = service.files().create(
        body={'name': filename, 'parents': [folder_id]},
        media_body=media,
        fields='id,webViewLink',
    ).execute()
    return result['webViewLink']


def build_on_template(letter_docx_path: str, filename: str) -> str:
    """Export a letter by appending its content to the downloaded Drive template.

    Flow:
    1. Download the Drive template.
    2. Open both documents (template and generated letter).
    3. Find the first content element in the generated letter (skip leading empty paragraphs).
    4. Append all subsequent content elements (paragraphs and tables) to the template body before its sectPr.
    5. Save, post-process fonts (David -> David Libre), and upload.
    """
    import tempfile, copy, re as _re, os as _os, zipfile as _zf
    from docx import Document
    from docx.oxml.ns import qn

    # ── 1. Download template ──────────────────────────────────────────
    tmp_tpl = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp_tpl.close()
    _download_drive_file(DRIVE_TEMPLATE_ID, tmp_tpl.name)

    # ── 2. Open both docs ─────────────────────────────────────────────
    tpl_doc = Document(tmp_tpl.name)
    ltr_doc = Document(letter_docx_path)
    tpl_body = tpl_doc.element.body
    ltr_body = ltr_doc.element.body

    # ── 3. Find the content elements in the generated letter ──────────
    # We want to skip the leading empty spacer paragraphs.
    P = qn('w:p')
    TBL = qn('w:tbl')
    T = qn('w:t')
    SEC = qn('w:sectPr')

    def _para_text(el):
        return ''.join(t.text or '' for t in el.iter(T))

    ltr_elements = [c for c in ltr_body if c.tag in (P, TBL)]
    
    first_content_idx = 0
    for i, el in enumerate(ltr_elements):
        if el.tag == TBL:
            first_content_idx = i
            break
        elif el.tag == P:
            # Check if paragraph has text or images/drawings
            if _para_text(el).strip() or el.findall('.//' + qn('w:drawing')):
                first_content_idx = i
                break
    else:
        first_content_idx = 0

    content_elements = ltr_elements[first_content_idx:]

    # ── 3b. Add signature to tpl_doc if it exists in ltr_doc ───────────
    SIGNATURE_PATH = _os.path.join(BASE_DIR, 'חתימה.png')
    sig_rId = None
    if _os.path.exists(SIGNATURE_PATH):
        try:
            sig_rId, _ = tpl_doc.part.get_or_add_image(SIGNATURE_PATH)
        except Exception:
            pass

    # ── 4. Append content elements to the template ────────────────────
    sectPr = tpl_body.find(SEC)
    for el in content_elements:
        copied_el = copy.deepcopy(el)
        
        # If we have a signature image registered, map any copied image relationships to it
        if sig_rId:
            blips = copied_el.findall('.//' + qn('a:blip'))
            for blip in blips:
                embed_attr = qn('r:embed')
                old_rId = blip.get(embed_attr)
                if old_rId and old_rId in ltr_doc.part.rels:
                    rel = ltr_doc.part.rels[old_rId]
                    if 'image' in rel.reltype:
                        blip.set(embed_attr, sig_rId)

        if sectPr is not None:
            sectPr.addprevious(copied_el)
        else:
            tpl_body.append(copied_el)

    # ── 5. Save to temp ───────────────────────────────────────────────
    tmp_res = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp_res.close()
    tpl_doc.save(tmp_res.name)

    # ── 6. Font substitution via ZIP post-process (David -> David Libre)
    with _zf.ZipFile(tmp_res.name, 'r') as zin:
        zfiles = {n: zin.read(n) for n in zin.namelist()}
    
    for doc_key in ('word/document.xml', 'word/styles.xml'):
        if doc_key in zfiles:
            xml = zfiles[doc_key].decode('utf-8')
            xml = _re.sub(r'(w:(?:ascii|hAnsi|cs))="David"', r'\1="David Libre"', xml)
            zfiles[doc_key] = xml.encode('utf-8')

    with _zf.ZipFile(tmp_res.name, 'w', _zf.ZIP_DEFLATED) as zout:
        for n, d in zfiles.items():
            zout.writestr(n, d)

    # ── 7. Upload ─────────────────────────────────────────────────────
    url = _upload_raw(tmp_res.name, filename)

    _os.unlink(tmp_tpl.name)
    _os.unlink(tmp_res.name)
    return url


def upload_file(file_path: str, filename: str) -> str:
    """Upload a file (PDF or DOCX) to Google Drive, return webViewLink.
    For DOCX files, strips the full-page header image so body content
    is fully visible in Google Drive's viewer."""
    import tempfile, os as _os
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload

    creds = _load_creds()
    if not creds:
        raise RuntimeError('אין הרשאות Google Drive — יש לבצע אימות מחדש')

    if filename.lower().endswith('.docx'):
        return build_on_template(file_path, filename)

    if filename.lower().endswith('.pdf'):
        mimetype = 'application/pdf'
    else:
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    service = build('drive', 'v3', credentials=creds)
    folder_id = _get_target_folder(service)

    media = MediaFileUpload(file_path, mimetype=mimetype, resumable=False)
    result = service.files().create(
        body={'name': filename, 'parents': [folder_id]},
        media_body=media,
        fields='id,webViewLink',
    ).execute()

    return result['webViewLink']


# backward compat
def upload_docx(file_path: str, filename: str) -> str:
    return upload_file(file_path, filename)
